"""
Memoria de cálculo Word — EnginePro Losas · CRÉER Ingeniería
Versión mejorada: espaciado generoso, marca de agua, encabezados/pies de página,
sección transversal premium con diferenciación de colores.
"""
import os
import sys
import datetime
import tempfile
import math
import numpy as np

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, Polygon, Circle, Rectangle
from matplotlib.figure import Figure
from matplotlib import gridspec

from docx import Document as DocxDocument
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from .motor import VERSION, EMPRESA, MALLAS, GRAFILES, MotorEstructural

# Colores corporativos CRÉER
C_PRIMARY   = "#1A1A2E"
C_ACCENT    = "#E03030"
C_INF       = "#C0392B"   # rojo barras inferiores
C_SUP       = "#1A5276"   # azul barras superiores
C_CONC      = "#B8C0CC"
C_DIM       = "#4A5568"
C_GRID      = "#E8ECF2"


class GraficasExport:
    """Gráficas matplotlib para insertar en Word — estilo 2026."""

    @staticmethod
    def setup_style():
        plt.rcParams.update({
            'font.family': 'sans-serif',
            'font.sans-serif': ['Arial', 'Helvetica', 'sans-serif'],
            'font.size': 9,
            'axes.titlesize': 11,
            'axes.titleweight': 'bold',
            'axes.titlecolor': C_PRIMARY,
            'axes.labelsize': 9,
            'axes.labelcolor': '#333',
            'axes.grid': True,
            'grid.alpha': 0.4,
            'grid.linestyle': '--',
            'grid.color': '#CBD5E1',
            'figure.facecolor': '#ffffff',
            'axes.facecolor': '#F8FAFC',
            'axes.edgecolor': '#CBD5E1',
            'axes.linewidth': 0.8,
            'xtick.color': '#555',
            'ytick.color': '#555',
        })

    @staticmethod
    def anotar_extremo(ax, x, y_arr, color, fmt='.2f', units='', top=True):
        idx = np.argmax(y_arr) if top else np.argmin(y_arr)
        val = y_arr[idx]
        if abs(val) < 1e-6: return
        ax.annotate(
            f'{val:{fmt}} {units}',
            xy=(x[idx], val), fontsize=8, fontweight='bold', color=color,
            ha='center', va='bottom' if top else 'top',
            xytext=(0, 14 if top else -14), textcoords='offset points',
            bbox=dict(boxstyle='round,pad=0.3', facecolor='white', edgecolor=color, alpha=0.92, linewidth=1.2),
            arrowprops=dict(arrowstyle='->', color=color, lw=0.9)
        )

    @staticmethod
    def _rebar_label(malla, grafil):
        lbl = malla
        if grafil and grafil != "Sin Grafil":
            lbl += " + " + grafil
        return lbl

    @staticmethod
    def crear_figura_principal(R):
        GraficasExport.setup_style()
        fig = plt.figure(figsize=(11, 15))
        gs_main = gridspec.GridSpec(4, 1, height_ratios=[0.75, 1.3, 1.0, 1.05],
                                    hspace=0.48, top=0.96, bottom=0.04, left=0.10, right=0.96)
        x = np.array(R.x_global)
        L_total = sum(R.L_list)
        apoyos_x = [0] + list(np.cumsum(R.L_list))

        # ─ 1. ESQUEMA ESTRUCTURAL ─
        ax1 = fig.add_subplot(gs_main[0])
        ax1.set_xlim(-0.18*L_total, L_total*1.18)
        ax1.set_ylim(-0.7, 1.8)
        ax1.set_aspect('equal')
        ax1.axis('off')
        ax1.set_title('ESQUEMA ESTRUCTURAL', fontsize=12, fontweight='bold',
                      color=C_PRIMARY, pad=8)

        # Losa
        ax1.plot([0, L_total], [0, 0], color=C_PRIMARY, lw=6, solid_capstyle='round', zorder=3)
        ax1.fill_between([0, L_total], [0,0], [-0.05,-0.05], color=C_CONC, alpha=0.4)

        # Apoyos triángulos
        for xp in apoyos_x:
            tri = Polygon([[xp,0],[xp-0.07,-0.22],[xp+0.07,-0.22]],
                          closed=True, facecolor='#334155', edgecolor='#1A2A3A', lw=1.2, zorder=4)
            ax1.add_patch(tri)
            ax1.plot([xp-0.12,-0.12+xp+0.24],[xp*0-0.28,xp*0-0.28], color='#334155', lw=2)

        # Cargas distribuidas (flechas)
        n_arrows = max(3, int(L_total * 2))
        for i in range(n_arrows):
            xa = L_total / n_arrows * (i + 0.5)
            ax1.annotate('', xy=(xa, 0), xytext=(xa, 0.55),
                         arrowprops=dict(arrowstyle='->', color=C_ACCENT, lw=1.3))
        ax1.plot([0, L_total], [0.6, 0.6], color=C_ACCENT, lw=1.8)
        ax1.text(L_total/2, 0.75, f'Wu = {R.wu_max:.2f} kN/m', ha='center',
                 fontsize=9, color=C_ACCENT, fontweight='bold')

        # Etiquetas vanos
        for i, L in enumerate(R.L_list):
            xmid = apoyos_x[i] + L/2
            ax1.text(xmid, -0.45, f'L{i+1} = {L:.2f} m', ha='center', fontsize=9,
                     fontweight='bold', color=C_PRIMARY)
            # cotas
            ax1.annotate('', xy=(apoyos_x[i+1], -0.58), xytext=(apoyos_x[i], -0.58),
                         arrowprops=dict(arrowstyle='<->', color=C_DIM, lw=0.9))

        # ─ 2. MOMENTOS ─
        ax2 = fig.add_subplot(gs_main[1])
        M_pos = np.array(R.M_env_max)
        M_neg = np.array(R.M_env_min)
        ax2.fill_between(x, M_pos, 0, where=M_pos>0, color='#EF4444', alpha=0.15, label='M+ (tracción inf)')
        ax2.fill_between(x, M_neg, 0, where=M_neg<0, color='#3B82F6', alpha=0.15, label='M- (tracción sup)')
        ax2.plot(x, M_pos, color='#EF4444', lw=2.2, zorder=5)
        ax2.plot(x, M_neg, color='#3B82F6', lw=2.2, zorder=5)
        ax2.axhline(0, color='#334155', lw=0.7)
        GraficasExport.anotar_extremo(ax2, x, M_pos, '#EF4444', units='kN·m')
        GraficasExport.anotar_extremo(ax2, x, M_neg, '#3B82F6', units='kN·m', top=False)
        for xp in apoyos_x:
            ax2.axvline(xp, color='#94A3B8', lw=0.9, ls='--', alpha=0.6)
        ax2.set_title('ENVOLVENTE DE MOMENTOS FLECTORES', fontsize=11, fontweight='bold', color=C_PRIMARY)
        ax2.set_ylabel('Mu [kN·m]', fontsize=9)
        ax2.invert_yaxis()
        ax2.legend(loc='lower right', fontsize=8, framealpha=0.9)

        # ─ 3. CORTANTES ─
        ax3 = fig.add_subplot(gs_main[2])
        V_pos = np.array(R.V_env_max)
        V_neg = np.array(R.V_env_min)
        ax3.fill_between(x, V_pos, 0, where=V_pos>0, color='#22C55E', alpha=0.15)
        ax3.fill_between(x, V_neg, 0, where=V_neg<0, color='#22C55E', alpha=0.10)
        ax3.plot(x, V_pos, color='#22C55E', lw=2.2, label='Vu máx')
        ax3.plot(x, V_neg, color='#22C55E', lw=2.2, ls='--', label='Vu mín')
        ax3.axhline(R.phi_vc, color='#F59E0B', lw=1.5, ls='-.', label=f'φVc = {R.phi_vc:.2f} kN')
        ax3.axhline(-R.phi_vc, color='#F59E0B', lw=1.5, ls='-.')
        ax3.axhline(0, color='#334155', lw=0.7)
        for xp in apoyos_x:
            ax3.axvline(xp, color='#94A3B8', lw=0.9, ls='--', alpha=0.6)
        ax3.set_title('ENVOLVENTE DE CORTANTES', fontsize=11, fontweight='bold', color=C_PRIMARY)
        ax3.set_ylabel('Vu [kN]', fontsize=9)
        ax3.legend(loc='upper right', fontsize=8, framealpha=0.9)

        # ─ 4. DEFLEXIONES ─
        ax4 = fig.add_subplot(gs_main[3])
        has_data = False
        if R.delta_DL_x is not None and len(R.delta_DL_x) == len(x):
            ax4.plot(x, np.array(R.delta_DL_x)*1000, color='#A78BFA', lw=1.8, ls='--',
                     label=f'δ(D+L) = {R.delta_DL_mm:.2f} mm')
            has_data = True
        if R.delta_LP_x is not None and len(R.delta_LP_x) == len(x):
            ax4.fill_between(x, np.array(R.delta_LP_x)*1000, 0, alpha=0.12, color='#EF4444')
            lp_ratio = int(sum(R.L_list)*1000/R.delta_LP_mm) if R.delta_LP_mm > 1e-6 else 0
            lp_label = f'δLP = {R.delta_LP_mm:.2f} mm  (L/{lp_ratio})' if lp_ratio > 0 else f'δLP = {R.delta_LP_mm:.2f} mm'
            ax4.plot(x, np.array(R.delta_LP_x)*1000, color='#EF4444', lw=2.5, label=lp_label)
            has_data = True
        if has_data:
            ax4.invert_yaxis()
        ax4.axhline(0, color='#334155', lw=0.7)
        for xp in apoyos_x:
            ax4.axvline(xp, color='#94A3B8', lw=0.9, ls='--', alpha=0.6)
        ax4.set_title('DEFLEXIONES', fontsize=11, fontweight='bold', color=C_PRIMARY)
        ax4.set_ylabel('δ [mm]', fontsize=9)
        ax4.set_xlabel('Longitud [m]', fontsize=9)
        if has_data:
            ax4.legend(loc='lower right', fontsize=8, framealpha=0.9)

        fig.suptitle(f'{EMPRESA} — EnginePro Losas v{VERSION}',
                     fontsize=8.5, color='#94A3B8', y=0.995)
        return fig

    @staticmethod
    def crear_figura_seccion(R):
        """Sección transversal — escala uniforme real."""
        GraficasExport.setup_style()

        # ── Escala uniforme (mismo criterio que el SVG del browser) ──
        # 1 unidad = 10cm. Losa mínimo 0.8u = 8cm visual, máximo 2.0u = 20cm visual.
        U_MIN, U_MAX = 0.8, 2.0
        h_natural = R.h / 10.0          # unidades naturales
        if h_natural < U_MIN:
            scale = U_MIN / h_natural    # escalar todo
        elif h_natural > U_MAX:
            scale = U_MAX / h_natural
        else:
            scale = 1.0

        b   = 10.0 * scale              # 100cm a escala
        h   = R.h  / 10.0 * scale      # altura losa a escala
        rec = 0.30 * scale              # 3cm recubrimiento a escala

        # Radios de barras — escala real, sin factor de distorsión
        # r_mm_to_u: mm → unidades de dibujo (1u=10cm → 1u=100mm)
        def bar_r(db_mm, sp_u):
            r_real = (db_mm / 2.0) / 100.0 * scale
            r_max  = sp_u * 0.42
            return min(r_real, r_max)

        db_mi = MALLAS[R.malla_inf]["diametro"]
        db_ms = MALLAS[R.malla_sup]["diametro"]
        db_gi = GRAFILES[R.grafil_inf]["diametro"]
        db_gs = GRAFILES[R.grafil_sup]["diametro"]

        sep_inf_u = MALLAS[R.malla_inf]["sep"] / 10.0 * scale
        sep_sup_u = MALLAS[R.malla_sup]["sep"] / 10.0 * scale
        n_bars_inf = max(2, min(int(b / max(sep_inf_u, 0.05)), 16))
        n_bars_sup = max(2, min(int(b / max(sep_sup_u, 0.05)), 16))
        sp_inf = b / (n_bars_inf + 1)
        sp_sup = b / (n_bars_sup + 1)

        r_mi = bar_r(db_mi, sp_inf)
        r_ms = bar_r(db_ms, sp_sup)
        r_gi = bar_r(db_gi, sp_inf) if db_gi > 0 else 0
        r_gs = bar_r(db_gs, sp_sup) if db_gs > 0 else 0

        # figsize proporcional al rango de ejes para que set_aspect('equal') no distorsione
        x_margin, y_top, y_bot = 4.2 * scale, 1.2 * scale, 2.2 * scale
        xlim = (-1.0 * scale, b + x_margin)
        ylim = (-y_bot, h + y_top)
        x_range = xlim[1] - xlim[0]
        y_range = ylim[1] - ylim[0]
        fig_w = 10.0
        fig_h = max(2.5, min(8.0, fig_w * y_range / x_range))

        fig, ax = plt.subplots(1, 1, figsize=(fig_w, fig_h))
        ax.set_aspect('equal')
        ax.axis('off')

        # Concreto
        rect = FancyBboxPatch((0, 0), b, h, boxstyle="round,pad=0.04",
                              linewidth=2, edgecolor='#2C3E50', facecolor='#C8CDD8', zorder=1)
        ax.add_patch(rect)
        for i in range(int(b * 6 / scale)):
            xi = i * 0.18 * scale
            ax.plot([xi, min(xi+h, b)], [0, min(h, b-xi)],
                    color='#5A6980', lw=0.15, alpha=0.4, zorder=2)

        # Líneas guía
        ax.plot([0.2*scale, b-0.2*scale], [rec, rec],      color='#3B82F6', lw=0.5, ls='--', alpha=0.3, zorder=3)
        ax.plot([0.2*scale, b-0.2*scale], [h-rec, h-rec],  color='#EF4444', lw=0.5, ls='--', alpha=0.3, zorder=3)

        # Barras INF (rojo)
        for i in range(n_bars_inf):
            bx = (i+1)*sp_inf
            ax.add_patch(Circle((bx, rec), r_mi, color='#EF4444', ec='#7F1D1D', lw=1.0, zorder=6))
        if r_gi > 0:
            for i in range(n_bars_inf-1):
                bx = (i+1)*sp_inf + sp_inf/2
                ax.add_patch(Circle((bx, rec), r_gi, color='#F87171', ec='#7F1D1D', lw=0.8, alpha=0.9, zorder=6))

        # Barras SUP (azul)
        for i in range(n_bars_sup):
            bx = (i+1)*sp_sup
            ax.add_patch(Circle((bx, h-rec), r_ms, color='#3B82F6', ec='#1E3A5F', lw=1.0, zorder=6))
        if r_gs > 0:
            for i in range(n_bars_sup-1):
                bx = (i+1)*sp_sup + sp_sup/2
                ax.add_patch(Circle((bx, h-rec), r_gs, color='#60A5FA', ec='#1E3A5F', lw=0.8, alpha=0.9, zorder=6))

        # Recubrimiento
        ax.annotate('', xy=(0.25*scale, rec), xytext=(0.25*scale, 0),
                    arrowprops=dict(arrowstyle='<->', color=C_DIM, lw=0.9))
        ax.text(0.45*scale, rec/2, f'r={int(rec/scale*10):.0f}cm', fontsize=7, color=C_DIM, va='center')

        # Cotas
        def dim_h(x1, x2, y, label):
            ax.annotate('', xy=(x2,y), xytext=(x1,y), arrowprops=dict(arrowstyle='<->', color=C_DIM, lw=0.9))
            for xi in [x1, x2]: ax.plot([xi,xi],[y-0.1*scale,y+0.1*scale], color=C_DIM, lw=0.7)
            ax.text((x1+x2)/2, y-0.25*scale, label, ha='center', va='top', fontsize=8, color=C_DIM, fontweight='bold')

        def dim_v(x, y1, y2, label):
            ax.annotate('', xy=(x,y2), xytext=(x,y1), arrowprops=dict(arrowstyle='<->', color=C_DIM, lw=0.9))
            for yi in [y1, y2]: ax.plot([x-0.1*scale,x+0.1*scale],[yi,yi], color=C_DIM, lw=0.7)
            ax.text(x+0.18*scale, (y1+y2)/2, label, ha='left', va='center', fontsize=8, color=C_DIM, fontweight='bold')

        dim_v(b+0.7*scale, 0, h,     f'h = {R.h:.0f} cm')
        dim_v(b+2.1*scale, 0, h-rec, f'd = {R.d:.0f} cm')
        dim_h(0, b, -1.3*scale, 'b = 100 cm (franja unitaria)')

        # Etiquetas con fondo
        lbl_sup = GraficasExport._rebar_label(R.malla_sup, R.grafil_sup)
        lbl_inf = GraficasExport._rebar_label(R.malla_inf, R.grafil_inf)
        ax.text(b/2, h+0.55*scale, f'SUPERIOR: {lbl_sup}', ha='center', fontsize=8.5,
                color='#1E3A8A', fontweight='bold',
                bbox=dict(boxstyle='round,pad=0.3', facecolor='#EFF6FF', edgecolor='#3B82F6', alpha=0.95))
        ax.text(b/2, -0.65*scale, f'INFERIOR: {lbl_inf}', ha='center', fontsize=8.5,
                color='#7F1D1D', fontweight='bold',
                bbox=dict(boxstyle='round,pad=0.3', facecolor='#FEF2F2', edgecolor='#EF4444', alpha=0.95))

        ax.set_xlim(xlim)
        ax.set_ylim(ylim)
        ax.set_title('SECCIÓN TRANSVERSAL Y ARMADO — Losa Maciza en 1 Dirección',
                     fontsize=11, fontweight='bold', color=C_PRIMARY, pad=10)

        # Leyenda — debajo de las etiquetas, fuera del dibujo
        legend_elements = [
            mpatches.Patch(facecolor='#EF4444', edgecolor='#7F1D1D',
                           label=f'INF: {R.malla_inf} — Ø{db_mi:.2f}mm c/{MALLAS[R.malla_inf]["sep"]}cm'),
            mpatches.Patch(facecolor='#3B82F6', edgecolor='#1E3A5F',
                           label=f'SUP: {R.malla_sup} — Ø{db_ms:.2f}mm c/{MALLAS[R.malla_sup]["sep"]}cm'),
            mpatches.Patch(facecolor='#C8CDD8', edgecolor='#2C3E50',
                           label=f"Concreto f'c = {R.fc:.0f} MPa"),
        ]
        if r_gi > 0:
            legend_elements.insert(1, mpatches.Patch(facecolor='#F87171', edgecolor='#7F1D1D',
                                                      label=f'Grafil INF Ø{db_gi:.1f}mm'))
        if r_gs > 0:
            legend_elements.insert(-1, mpatches.Patch(facecolor='#60A5FA', edgecolor='#1E3A5F',
                                                       label=f'Grafil SUP Ø{db_gs:.1f}mm'))
        ax.legend(handles=legend_elements, loc='upper left', fontsize=7.5,
                  framealpha=0.92, edgecolor='#CBD5E1',
                  bbox_to_anchor=(0, -0.04), bbox_transform=ax.transAxes, ncol=2)

        fig.tight_layout(pad=1.2)
        return fig



class MemoriaWord:
    """Genera la memoria de cálculo Word — estilo profesional 2026."""

    COL_PRIMARY = RGBColor(0x1A, 0x1A, 0x2E)
    COL_ACCENT  = RGBColor(0xE0, 0x30, 0x30)
    COL_OK      = RGBColor(0x15, 0x80, 0x3D)
    COL_FAIL    = RGBColor(0xC0, 0x20, 0x20)
    COL_GRAY    = RGBColor(0x64, 0x74, 0x8B)
    COL_DARK    = RGBColor(0x1E, 0x29, 0x3B)
    COL_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
    COL_BLUE    = RGBColor(0x1D, 0x4E, 0xD8)
    COL_LIGHT   = RGBColor(0xF8, 0xFA, 0xFC)

    def _shd(self, cell, hex_color):
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shd)

    def _table(self, doc, headers, rows, widths=None, header_color="1A1A2E"):
        n = len(headers)
        t = doc.add_table(rows=1+len(rows), cols=n)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.style = 'Table Grid'
        # Header
        for i, h in enumerate(headers):
            c = t.rows[0].cells[i]; c.text = ""
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(h); r.bold = True; r.font.size = Pt(9)
            r.font.color.rgb = self.COL_WHITE; r.font.name = 'Arial'
            self._shd(c, header_color)
        # Rows
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                c = t.rows[ri+1].cells[ci]; c.text = ""
                p = c.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
                s = str(val); r = p.add_run(s)
                r.font.size = Pt(9); r.font.name = 'Arial'
                if s == "CUMPLE":   r.bold = True; r.font.color.rgb = self.COL_OK
                elif s == "NO CUMPLE": r.bold = True; r.font.color.rgb = self.COL_FAIL
                else: r.font.color.rgb = self.COL_DARK
                if ri % 2 == 1: self._shd(c, "F0F4F8")
        if widths:
            for i, w in enumerate(widths):
                for row in t.rows:
                    row.cells[i].width = Cm(w)
        return t

    def _heading(self, doc, text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.color.rgb = self.COL_PRIMARY if level == 1 else self.COL_DARK
            run.font.name = 'Arial'
        if level == 1:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="12" w:space="1" w:color="E03030"/>'
                f'</w:pBdr>')
            pPr.insert(0, pBdr)
        return h

    def _para(self, doc, text, bold=False, color=None, size=10, align=None, italic=False, space_before=0, space_after=6):
        p = doc.add_paragraph()
        if align: p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.bold = bold; run.italic = italic
        run.font.size = Pt(size); run.font.name = 'Arial'
        run.font.color.rgb = color or self.COL_DARK
        return p

    def _eq(self, doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.2)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text); r.font.size = Pt(9.5); r.font.name = 'Courier New'
        r.font.color.rgb = self.COL_DARK
        return p

    def _ref(self, doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.2)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text); r.font.size = Pt(8.5); r.font.name = 'Arial'
        r.font.color.rgb = self.COL_GRAY; r.italic = True
        return p

    def _resultado(self, doc, texto_antes, valor, texto_despues, cumple, size=10):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(texto_antes + valor + texto_despues)
        r1.font.size = Pt(size); r1.font.name = 'Arial'
        estado = "CUMPLE ✔" if cumple else "NO CUMPLE ✘"
        r2 = p.add_run(f"  {estado}")
        r2.bold = True; r2.font.size = Pt(size); r2.font.name = 'Arial'
        r2.font.color.rgb = self.COL_OK if cumple else self.COL_FAIL
        return p

    def _watermark_page(self, section):
        """Agrega marca de agua CRÉER Ingeniería en encabezado."""
        # Header with brand
        header = section.header
        header.is_linked_to_previous = False
        ht = header.paragraphs[0]
        ht.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ht.paragraph_format.space_after = Pt(0)
        pPr = ht._p.get_or_add_pPr()
        pPr.insert(0, parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="6" w:space="4" w:color="E03030"/>'
            f'</w:pBdr>'))
        r = ht.add_run("CRÉER Ingeniería · EnginePro Losas")
        r.font.size = Pt(8); r.font.name = 'Arial'
        r.font.color.rgb = self.COL_GRAY; r.bold = True

    def _add_footer(self, section, fecha):
        footer = section.footer
        footer.is_linked_to_previous = False
        ft = footer.paragraphs[0]
        ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = ft._p.get_or_add_pPr()
        pPr.insert(0, parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="4" w:color="CBD5E1"/>'
            f'</w:pBdr>'))
        r = ft.add_run(f"Documento generado por EnginePro Losas v{VERSION} · CRÉER Ingeniería · {fecha}  —  creeringenieria.com")
        r.font.size = Pt(7.5); r.font.name = 'Arial'
        r.font.color.rgb = self.COL_GRAY

    def generar(self, R, filepath, datos_proyecto):
        doc = DocxDocument()

        nombre    = datos_proyecto.get("nombre", EMPRESA)
        matricula = datos_proyecto.get("matricula", "Profesional Especializado")
        proyecto  = datos_proyecto.get("proyecto", "Proyecto Estructural")
        ubicacion = datos_proyecto.get("ubicacion", "Colombia")
        fecha     = datos_proyecto.get("fecha", datetime.date.today().strftime("%d/%m/%Y"))

        # Márgenes
        for sec in doc.sections:
            sec.top_margin    = Cm(2.5)
            sec.bottom_margin = Cm(2.5)
            sec.left_margin   = Cm(2.8)
            sec.right_margin  = Cm(2.5)
            self._watermark_page(sec)
            self._add_footer(sec, fecha)

        # ══════════ PORTADA ══════════
        doc.add_paragraph()  # espacio

        # Línea roja superior
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(18)
        pPr = p._p.get_or_add_pPr()
        pPr.insert(0, parse_xml(
            f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="36" w:space="6" w:color="E03030"/></w:pBdr>'))

        # Empresa
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(nombre); r.bold = True; r.font.size = Pt(28)
        r.font.color.rgb = self.COL_PRIMARY; r.font.name = 'Arial'

        # Matrícula
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(32)
        r = p.add_run(matricula); r.font.size = Pt(13)
        r.font.color.rgb = self.COL_GRAY; r.italic = True; r.font.name = 'Arial'

        # Título
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run("MEMORIA DE CÁLCULO ESTRUCTURAL"); r.bold = True; r.font.size = Pt(24)
        r.font.color.rgb = self.COL_PRIMARY; r.font.name = 'Arial'

        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run("LOSA MACIZA EN UNA DIRECCIÓN"); r.bold = True; r.font.size = Pt(16)
        r.font.color.rgb = self.COL_ACCENT; r.font.name = 'Arial'

        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(28)
        r = p.add_run("Normativa: NSR-10 Título C / ACI 318"); r.font.size = Pt(12)
        r.font.color.rgb = self.COL_GRAY; r.font.name = 'Arial'

        # Tabla info proyecto
        self._table(doc, ["Campo", "Detalle"],
                    [["PROYECTO:", proyecto], ["UBICACIÓN:", ubicacion],
                     ["FECHA:", fecha], ["SOFTWARE:", f"EnginePro Losas v{VERSION} · {nombre}"]],
                    [4.5, 12.5])

        doc.add_paragraph().paragraph_format.space_after = Pt(16)

        # Línea roja final portada
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pPr.insert(0, parse_xml(
            f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="28" w:space="4" w:color="E03030"/></w:pBdr>'))

        doc.add_page_break()

        # ══════════ 1. CRITERIOS ══════════
        self._heading(doc, "1. CRITERIOS DE DISEÑO")

        doc.add_heading("1.1  Normativa de referencia", level=2)
        normas = [
            "NSR-10 Título B — Cargas",
            "NSR-10 Título C — Concreto Estructural (basado en ACI 318-19)",
            "NSR-10 C.7.12 — Refuerzo por retracción y temperatura",
            "NSR-10 C.8.11.2 — Disposiciones de carga alternada",
            "NSR-10 C.9.5 — Control de deflexiones",
            "NSR-10 C.10 — Flexión y carga axial",
            "NSR-10 C.11.2 — Resistencia al cortante",
        ]
        for n in normas:
            self._para(doc, f"  •  {n}", size=9.5, space_after=3)

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

        doc.add_heading("1.2  Materiales", level=2)
        self._table(doc, ["Parámetro", "Valor", "Norma / Referencia"],
                    [["Resistencia concreto f'c", f"{R.fc:.1f} MPa", "NSR-10 C.5"],
                     ["Fluencia acero fy", f"{R.fy:.1f} MPa", "NSR-10 C.3.5"],
                     ["Módulo elástico Ec", f"{R.Ec:.0f} MPa", "Ec = 4700·√(f'c) — C.8.5.1"],
                     ["Factor β₁", f"{R.beta1:.3f}", "NSR-10 C.10.2.7.3"],
                     ["Peso unitario concreto", "2 400 kgf/m³", "NSR-10 C.8.5"]],
                    [6, 3.5, 8.5])
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

        # ══════════ 2. GEOMETRÍA ══════════
        doc.add_page_break()
        self._heading(doc, "2. GEOMETRÍA DE LA SECCIÓN")

        luces_str = " + ".join([f"{L:.2f} m" for L in R.L_list])
        self._para(doc, f"Configuración: {len(R.L_list)} vano(s)  —  Luces totales: {luces_str}  =  {sum(R.L_list):.2f} m",
                   bold=True, space_after=10)

        self._table(doc, ["Parámetro", "Valor"],
                    [["Tipo de losa", "Maciza en 1 dirección"],
                     ["Espesor total h", f"{R.h:.1f} cm"],
                     ["Altura útil d", f"{R.d:.1f} cm"],
                     ["Recubrimiento (hasta centro acero)", "3.0 cm"],
                     ["Ancho de diseño b", "100 cm (franja unitaria 1 m)"]],
                    [6.5, 11.5])
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

        doc.add_heading("2.1  Verificación de espesor mínimo — Tabla NSR-10 C.9.5(a)", level=2)
        self._eq(doc, f"h_min = L / factor × (0.4 + fy/700) = {R.h_min_req:.2f} cm")
        self._resultado(doc, f"h provisto = {R.h:.1f} cm  ≥  h_min = {R.h_min_req:.2f} cm  →  ", "", "", R.cumple_h_min)
        self._ref(doc, "Ref: NSR-10 Tabla C.9.5(a) — Espesores mínimos para losas sin puntales")

        # ══════════ 3. CARGAS ══════════
        doc.add_page_break()
        self._heading(doc, "3. ANÁLISIS DE CARGAS")

        carga_rows = [
            ["Peso propio concreto (2400·h)", f"{R.pp_concreto:.1f} kgf/m²", "D"],
            ["Acabados + muros divisorios (CM adic)", f"{R.cm_adic:.1f} kgf/m²", "D"],
            ["CARGA MUERTA TOTAL (D)", f"{R.wd_serv:.1f} kgf/m²", "—"],
            ["CARGA VIVA (L)", f"{R.wl_serv:.1f} kgf/m²", "L"],
        ]
        self._table(doc, ["Carga", "Valor", "Tipo"], carga_rows, [8, 3.5, 2.5])
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

        doc.add_heading("3.1  Combinaciones de carga — NSR-10 B.2.4", level=2)
        self._eq(doc, f"U₁ = 1.2·D + 1.6·L = 1.2×{R.wd_kn:.3f} + 1.6×{R.wl_kn:.3f} = {R.wu_max:.3f} kN/m")
        self._eq(doc, f"U₂ = 1.4·D = 1.4×{R.wd_kn:.3f} = {1.4*R.wd_kn:.3f} kN/m")
        self._para(doc, f"Carga de diseño Wu = {R.wu_max:.3f} kN/m  (gobierna combinación U₁)", bold=True)
        self._ref(doc, "Se evaluaron patrones de carga alternada según C.8.11.2 para obtener la envolvente de diseño.")

        # ══════════ 4. ANÁLISIS ESTRUCTURAL ══════════
        self._heading(doc, "4. ANÁLISIS ESTRUCTURAL")
        self._para(doc, "Método: Tres Momentos (ecuación de Clapeyron) con envolventes de carga alternada NSR-10 C.8.11.2")
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

        patron_rows = [[f"{i+1}", nombre_p] for i, (nombre_p, _) in enumerate(R.patrones)]
        self._table(doc, ["#", "Patrón de Carga Alternada"], patron_rows, [2, 15])
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

        self._para(doc, f"Momento último positivo máximo (envolvente): Mu⁺ = {R.mu_pos:.3f} kN·m", bold=True, size=10.5)
        self._para(doc, f"Momento último negativo máximo (envolvente): Mu⁻ = {R.mu_neg:.3f} kN·m", bold=True, size=10.5)
        self._para(doc, f"Cortante último máximo a 'd' del apoyo: Vu = {R.vu_max:.3f} kN", bold=True, size=10.5)

        # ══════════ 5. REFUERZO ══════════
        doc.add_page_break()
        self._heading(doc, "5. REFUERZO PROVISTO")

        self._table(doc,
            ["Posición", "Malla electrosoldada", "Grafil adicional", "As total (cm²/m)", "db (mm)"],
            [["Superior (apoyos — negativo)", R.malla_sup, R.grafil_sup, f"{R.as_sup:.3f}", f"{MALLAS[R.malla_sup]['diametro']:.1f}"],
             ["Inferior (vanos — positivo)", R.malla_inf, R.grafil_inf, f"{R.as_inf:.3f}", f"{MALLAS[R.malla_inf]['diametro']:.1f}"]],
            [4.5, 4, 3.5, 3.5, 2.5])
        doc.add_paragraph().paragraph_format.space_after = Pt(8)
        self._para(doc, f"As mínimo por temperatura y retracción (C.7.12.2.1: 0.0018·b·h): {R.as_min_temp:.3f} cm²/m")
        self._ref(doc, "Para fy ≤ 420 MPa: As_min = 0.0018·b·h — NSR-10 C.7.12.2.1")

        # ══════════ 6. VERIFICACIONES ══════════
        self._heading(doc, "6. VERIFICACIONES DE DISEÑO")

        doc.add_heading("6.1  Flexión negativa (apoyos) — NSR-10 C.10", level=2)
        if len(R.L_list) > 1:
            self._para(doc, f"As superior provisto = {R.as_sup:.3f} cm²/m")
            self._eq(doc, f"a = As·fy / (0.85·f'c·b) = {R.as_sup:.3f}×{R.fy:.0f} / (0.85×{R.fc:.0f}×100)")
            self._resultado(doc, f"φMn⁻ = {R.phi_mn_neg:.3f} kN·m  ≥  Mu⁻ = {R.mu_neg:.3f} kN·m  →  ", "", "", R.cumple_neg)
            self._resultado(doc, f"εt = {R.et_neg:.5f}  ≥  0.005 (sección controlada por tracción)  →  ", "", "", R.et_neg >= 0.005)
        else:
            self._para(doc, "No aplica — vano único simplemente apoyado, sin momento negativo.", italic=True)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        doc.add_heading("6.2  Flexión positiva (vanos) — NSR-10 C.10", level=2)
        self._para(doc, f"As inferior provisto = {R.as_inf:.3f} cm²/m")
        self._resultado(doc, f"φMn⁺ = {R.phi_mn_pos:.3f} kN·m  ≥  Mu⁺ = {R.mu_pos:.3f} kN·m  →  ", "", "", R.cumple_pos)
        self._resultado(doc, f"εt = {R.et_pos:.5f}  ≥  0.005  →  ", "", "", R.et_pos >= 0.005)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        doc.add_heading("6.3  Cortante a d del apoyo — NSR-10 C.11.2", level=2)
        self._eq(doc, f"φVc = φ·0.17·√(f'c)·bw·d = 0.75×0.17×√{R.fc:.1f}×100×{R.d:.1f} = {R.phi_vc:.3f} kN")
        self._resultado(doc, f"φVc = {R.phi_vc:.3f} kN  ≥  Vu = {R.vu_max:.3f} kN  →  ", "", "", R.cumple_cortante)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        doc.add_heading("6.4  Cuantía máxima — NSR-10 C.10.3.3", level=2)
        self._eq(doc, f"ρ_max = 0.75·ρ_b = {R.rho_max:.6f}")
        self._resultado(doc, f"ρ provisto = {R.rho_provisto:.6f}  ≤  ρ_max = {R.rho_max:.6f}  →  ", "", "", R.cumple_rho_max)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        doc.add_heading("6.5  Fisuración — NSR-10 C.10.6.4", level=2)
        self._para(doc, f"Esfuerzo acero servicio positivo: fs⁺ = {R.fs_pos:.2f} MPa")
        self._resultado(doc, f"z⁺ = {R.z_pos:.0f}  ≤  31 000  →  ", "", "", R.cumple_fisura_pos)
        if len(R.L_list) > 1:
            self._para(doc, f"Esfuerzo acero servicio negativo: fs⁻ = {R.fs_neg:.2f} MPa")
            self._resultado(doc, f"z⁻ = {R.z_neg:.0f}  ≤  31 000  →  ", "", "", R.cumple_fisura_neg)

        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        doc.add_heading("6.6  Deflexiones — NSR-10 C.9.5", level=2)
        self._table(doc, ["Parámetro", "Valor"],
                    [["Inercia bruta Ig", f"{R.Ig:.4e} m⁴/m"],
                     ["Inercia fisurada Icr", f"{R.Icr:.4e} m⁴/m"],
                     ["Momento de fisuración Mcr", f"{R.Mcr:.3f} kN·m/m"],
                     ["Factor λΔ (largo plazo)", f"{R.lambda_lp:.2f}"],
                     ["Deflexión inmediata D", f"{R.delta_D_mm:.3f} mm"],
                     ["Deflexión inmediata D+L", f"{R.delta_DL_mm:.3f} mm"]],
                    [7, 11])
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        self._resultado(doc, f"Deflexión solo viva: δL = {R.delta_L_mm:.3f} mm  ≤  L/360 = {R.perm_L:.3f} mm  →  ", "", "", R.cumple_delta_L)
        self._resultado(doc, f"Deflexión largo plazo: δLP = {R.delta_LP_mm:.3f} mm  ≤  L/480 = {R.perm_LP:.3f} mm  →  ", "", "", R.cumple_delta_LP)

        # ══════════ 7. TABLA RESUMEN ══════════
        doc.add_page_break()
        self._heading(doc, "7. TABLA RESUMEN DE VERIFICACIONES NSR-10")
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

        verif_rows = [
            ["Flexión positiva (φMn ≥ Mu⁺)", f"{R.phi_mn_pos:.3f} kN·m", f"{R.mu_pos:.3f} kN·m", "CUMPLE" if R.cumple_pos else "NO CUMPLE"],
            ["Flexión negativa (φMn ≥ Mu⁻)", f"{R.phi_mn_neg:.3f} kN·m", f"{R.mu_neg:.3f} kN·m", "CUMPLE" if R.cumple_neg else "NO CUMPLE"],
            ["Cortante a d (φVc ≥ Vu)", f"{R.phi_vc:.3f} kN", f"{R.vu_max:.3f} kN", "CUMPLE" if R.cumple_cortante else "NO CUMPLE"],
            ["Cuantía máxima (ρ ≤ 0.75·ρb)", f"{R.rho_provisto:.6f}", f"{R.rho_max:.6f}", "CUMPLE" if R.cumple_rho_max else "NO CUMPLE"],
            ["Fisuración + (z ≤ 31000)", f"{R.z_pos:.0f}", "31 000", "CUMPLE" if R.cumple_fisura_pos else "NO CUMPLE"],
            ["Fisuración − (z ≤ 31000)", f"{R.z_neg:.0f}", "31 000", "CUMPLE" if R.cumple_fisura_neg else "NO CUMPLE"],
            ["Deflexión viva (L/360)", f"{R.delta_L_mm:.3f} mm", f"{R.perm_L:.3f} mm", "CUMPLE" if R.cumple_delta_L else "NO CUMPLE"],
            ["Deflexión LP (L/480)", f"{R.delta_LP_mm:.3f} mm", f"{R.perm_LP:.3f} mm", "CUMPLE" if R.cumple_delta_LP else "NO CUMPLE"],
            ["As mínimo superior", f"{R.as_sup:.3f} cm²/m", f"{R.as_min_temp:.3f} cm²/m", "CUMPLE" if R.cumple_as_min_sup else "NO CUMPLE"],
            ["As mínimo inferior", f"{R.as_inf:.3f} cm²/m", f"{R.as_min_temp:.3f} cm²/m", "CUMPLE" if R.cumple_as_min_inf else "NO CUMPLE"],
            ["Espesor mínimo", f"{R.h:.1f} cm", f"{R.h_min_req:.2f} cm", "CUMPLE" if R.cumple_h_min else "NO CUMPLE"],
        ]
        self._table(doc, ["Verificación", "Calculado", "Permisible", "Estado"],
                    verif_rows, [6.5, 3.5, 3.5, 3.5])

        # Dictamen
        doc.add_paragraph().paragraph_format.space_after = Pt(16)
        color_hex = "158042" if R.estado == "CUMPLE" else "C02020"
        color_rgb = self.COL_OK if R.estado == "CUMPLE" else self.COL_FAIL
        icon = "✔" if R.estado == "CUMPLE" else "✘"
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        pPr = p._p.get_or_add_pPr()
        pPr.insert(0, parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="18" w:space="6" w:color="{color_hex}"/>'
            f'<w:bottom w:val="single" w:sz="18" w:space="6" w:color="{color_hex}"/>'
            f'</w:pBdr>'))
        r = p.add_run(f" {icon}  DICTAMEN FINAL: EL DISEÑO {R.estado} TODAS LAS VERIFICACIONES NSR-10  {icon} ")
        r.bold = True; r.font.size = Pt(15); r.font.name = 'Arial'; r.font.color.rgb = color_rgb

        # ══════════ 8. DIAGRAMAS ══════════
        doc.add_page_break()
        self._heading(doc, "8. DIAGRAMAS DE SOLICITACIONES")
        self._para(doc, "Los diagramas siguientes corresponden a la envolvente de diseño obtenida mediante el método de los Tres Momentos (Clapeyron) con patrones de carga alternada según NSR-10 C.8.11.2.", size=9.5, space_after=10)

        tmpdir = tempfile.gettempdir()
        _uid = os.getpid()
        try:
            import threading
            _uid = f"{os.getpid()}_{threading.get_ident()}"
        except Exception:
            pass
        path_diag = os.path.join(tmpdir, f"creer_diag_{_uid}.png")
        path_sec  = os.path.join(tmpdir, f"creer_sec_{_uid}.png")

        try:
            fig_main = GraficasExport.crear_figura_principal(R)
            # NO tight_layout(): GridSpec ya tiene top/bottom/left/right explícitos
            # + suptitle(y=0.995) — tight_layout() conflicta y lanza excepción
            fig_main.savefig(path_diag, dpi=220, facecolor='white')
            plt.close(fig_main)
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(path_diag, width=Cm(16.5))
        except Exception as e:
            import traceback; traceback.print_exc()
            self._para(doc, f"[Error generando diagramas: {e}]", italic=True, color=self.COL_GRAY)

        # ══════════ 9. SECCIÓN TRANSVERSAL ══════════
        doc.add_page_break()
        self._heading(doc, "9. SECCIÓN TRANSVERSAL Y ARMADO")
        self._para(doc, "Sección a escala representativa. Los diámetros de barras se muestran proporcionales a los diámetros reales. Azul = refuerzo superior (momentos negativos). Rojo = refuerzo inferior (momentos positivos).", size=9.5, space_after=10)

        try:
            fig_sec = GraficasExport.crear_figura_seccion(R)
            # NO tight_layout(): ya se llama dentro de crear_figura_seccion(pad=1.2)
            fig_sec.savefig(path_sec, dpi=220, facecolor='white')
            plt.close(fig_sec)
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(path_sec, width=Cm(15))
        except Exception as e:
            import traceback; traceback.print_exc()
            self._para(doc, f"[Error generando sección: {e}]", italic=True, color=self.COL_GRAY)

        # Pie final
        doc.add_paragraph()
        self._para(doc, f"Documento generado el {fecha} por EnginePro Losas v{VERSION} — {nombre}",
                   size=8.5, color=self.COL_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)

        doc.save(filepath)

        for pf in [path_diag, path_sec]:
            try:
                if os.path.exists(pf): os.remove(pf)
            except Exception: pass
